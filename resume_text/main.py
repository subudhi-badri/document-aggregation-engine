# ==============================================================================
#  resume_text/main.py - OFFLINE VERSION WITH OLLAMA
# ==============================================================================

import os
import json
import pytesseract
from PIL import Image
import docx
from io import BytesIO
from pdf2image import convert_from_path
import google.generativeai as genai  # Import the new library
from dotenv import load_dotenv
import pdfplumber # Import the new library

# --- Load Environment Variables ---
load_dotenv()

# --- Direct Google AI Client Configuration ---
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

# Ensure these paths are correct for your system
POPPLER_PATH = r"C:\Users\ASUS\Desktop\restart\document_aggregation\poppler-24.08.0\Library\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\ASUS\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

# This is the detailed instruction for the AI model
system_instruction_prompt = '''You are an expert resume parser. Your sole function is to extract structured information from the provided resume text and return it ONLY as a valid JSON object. Adhere strictly to the following format. Pay special attention to finding the LinkedIn profile URL for the "linkedin_id" field AND the GitHub username for the "github_id" field. For any fields where information is not found, use an empty string "" or an empty list []. Do not use null.
{
    "name": "string",
    "linkedin_id": "string",
    "github_id": "string",
    "leetcode_id": "string",
    "kaggle_id": "string",
    "contact_info": {"email": "string", "phone": "string", "location": "string"},
    "summary": "string",
    "experience": [{"title": "string", "company": "string", "dates": "string", "description": "string", "location": "string"}],
    "education": [{"school": "string", "degree": "string", "field_of_study": "string", "dates": "string", "grades": "string"}],
    "skills": ["string"],
    "certifications": [{"name": "string", "issuing_authority": "string", "dates": "string"}],
    "projects": [{"title": "string", "description": "string", "technologies": ["string"]}],
    "languages": ["string"],
    "achievements": ["string"]
}
'''

# --- HELPER FUNCTIONS FOR TEXT EXTRACTION (These do not need to be changed) ---

def extract_text_from_image(image):
    return pytesseract.image_to_string(image).strip()

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_PATH)
    except Exception as e:
        print("[!] 300 DPI failed, retrying without dpi...")
        try:
            images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)
        except Exception as e2:
            print("[!] PDF to image failed:", e2)
            return ""
    return text.join([extract_text_from_image(img) for img in images])

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img = Image.open(BytesIO(img_data))
                text += "\n" + extract_text_from_image(img)
            except Exception as e:
                print("[!] DOCX image OCR failed:", e)
    return text.strip()

# --- NEW: Structure-Aware Text & Link Extractor ---
def extract_text_and_links_with_pdfplumber(file_path):
    """
    Extracts visible text AND hyperlink URLs from a PDF.
    This is the primary, preferred method.
    """
    all_text = ""
    print("[*] Attempting structured extraction with pdfplumber...")
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract the visible text from the page
                page_text = page.extract_text()
                if page_text:
                    all_text += page_text + "\n"
                
                # Extract all hyperlink URLs on the page
                if page.hyperlinks:
                    for link in page.hyperlinks:
                        # Append the hidden URL to our text block so the AI can see it
                        if 'uri' in link:
                            all_text += f" [Link found: {link['uri']}] "
                            
        if all_text:
            print("[✓] pdfplumber successfully extracted text and hyperlinks.")
            return all_text
        else:
            print("[!] pdfplumber found no digital text. This might be a scanned document.")
            return None
    except Exception as e:
        print(f"[!] An error occurred with pdfplumber: {e}")
        return None

# --- EXISTING: OCR-Based Text Extractor (Our Fallback) ---
def extract_text_with_ocr(file_path):
    """
    Extracts text from a PDF using OCR. Use this as a fallback
    for scanned documents.
    """
    print("[*] Falling back to OCR extraction with Tesseract...")
    try:
        images = convert_from_path(file_path, poppler_path=POPPLER_PATH)
        full_text = ""
        for img in images:
            full_text += pytesseract.image_to_string(img) + "\n"
        print("[✓] Tesseract OCR extraction completed.")
        return full_text
    except Exception as e:
        print(f"[!] An error occurred during OCR extraction: {e}")
        return ""

# --- GOOGLE GEMINI JSON GENERATOR FUNCTION ---
def generate_json_with_google(resume_text):
    """
    Takes raw resume text and uses the direct Google Gemini API 
    to parse it into structured JSON.
    """
    generation_config = {
        "response_mime_type": "application/json",
    }
    model = genai.GenerativeModel(
        "gemini-1.5-flash-latest",
        generation_config=generation_config
    )
    user_prompt = f"{system_instruction_prompt}\n\n--- RESUME CONTENT TO PARSE ---\n{resume_text}"
    try:
        print("[*] Calling direct Google Gemini API to parse resume...")
        response = model.generate_content(user_prompt)
        ai_response_str = response.text
        print("[*] --- Raw AI Response Received from Google ---")
        print(ai_response_str)
        print("[*] --- End of Raw AI Response ---")
        if not ai_response_str:
            print("[!] Google AI returned an empty response.")
            return None
        return json.loads(ai_response_str)
    except Exception as e:
        with open("resume_text/error.log", "a", encoding="utf-8") as logf:
            logf.write(f"[!] Error processing Google Gemini API response: {e}\n")
        print(f"[!] An error occurred while processing the Google Gemini API response: {e}")
        return None

# --- MAIN FUNCTION ---
def resume_extractor(file_path):
    """
    Main function to extract resume data using a hybrid strategy.
    Tries structured extraction first, then falls back to OCR.
    """
    filename = os.path.basename(file_path)
    # Determine file type
    if filename.endswith(".pdf"):
        # 1. Try the primary method (pdfplumber) first
        resume_text = extract_text_and_links_with_pdfplumber(file_path)

        # 2. If it fails (returns None), use the fallback method (OCR)
        if not resume_text:
            resume_text = extract_text_with_ocr(file_path)

    elif filename.endswith(".docx"):
        # Existing docx logic
        doc = docx.Document(file_path)
        resume_text = "\n".join([para.text for para in doc.paragraphs])
    else:
        # Unsupported file type
        return None

    if not resume_text:
        print(f"[!] Failed to extract any text from {filename}.")
        return None

    # 3. Send the extracted text (which now includes hyperlinks) to the AI
    final_json = generate_json_with_google(resume_text)

    if final_json:
        # Save the structured JSON to a file
        output_dir = "resume_text/output"
        os.makedirs(output_dir, exist_ok=True)
        json_filename = os.path.splitext(filename)[0] + '.json'
        json_filepath = os.path.join(output_dir, json_filename)
        with open(json_filepath, 'w') as f:
            json.dump(final_json, f, indent=4)
        print(f"[✓] Extracted JSON saved to {json_filepath}")
        return final_json
    else:
        print("[!] Failed to generate structured JSON from the extracted text.")
        return None